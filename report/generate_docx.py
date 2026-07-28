import os
import docx
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

def create_report():
    doc = Document()

    # Page setup - Margins 1 inch
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)

    # Color Palette - Professional Navy Theme
    NAVY = RGBColor(30, 58, 138)       # #1E3A8A - Heading 1
    SLATE = RGBColor(51, 65, 85)       # #334155 - Heading 2
    DARK_TEXT = RGBColor(30, 41, 59)   # #1E293B - Body Text
    MUTED_TEXT = RGBColor(100, 116, 139) # #64748B - Subtitles/Footers

    # Helper XML functions
    def set_cell_background(cell, fill_hex):
        tcPr = cell._element.get_or_add_tcPr()
        shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
        tcPr.append(shd)

    def set_cell_margins(cell, top=140, bottom=140, left=180, right=180):
        tcPr = cell._element.get_or_add_tcPr()
        tcMar = parse_xml(f'<w:tcMar {nsdecls("w")}><w:top w:w="{top}" w:type="dxa"/><w:bottom w:w="{bottom}" w:type="dxa"/><w:left w:w="{left}" w:type="dxa"/><w:right w:w="{right}" w:type="dxa"/></w:tcMar>')
        tcPr.append(tcMar)

    def style_heading_1(p, text):
        p.paragraph_format.space_before = Pt(16)
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.keep_with_next = True
        run = p.add_run(text)
        run.font.name = 'Calibri'
        run.font.size = Pt(15)
        run.font.bold = True
        run.font.color.rgb = NAVY
        return run

    def style_heading_2(p, text):
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.keep_with_next = True
        run = p.add_run(text)
        run.font.name = 'Calibri'
        run.font.size = Pt(12.5)
        run.font.bold = True
        run.font.color.rgb = SLATE
        return run

    def add_body_p(doc, text="", bold_prefix="", space_after=4):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(space_after)
        p.paragraph_format.line_spacing = 1.15
        if bold_prefix:
            r_bold = p.add_run(bold_prefix)
            r_bold.font.name = 'Calibri'
            r_bold.font.size = Pt(11)
            r_bold.font.bold = True
            r_bold.font.color.rgb = DARK_TEXT
        if text:
            r_text = p.add_run(text)
            r_text.font.name = 'Calibri'
            r_text.font.size = Pt(11)
            r_text.font.color.rgb = DARK_TEXT
        return p

    def add_bullet_p(doc, text, bold_prefix=""):
        p = doc.add_paragraph(style='List Bullet')
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(3)
        p.paragraph_format.line_spacing = 1.15
        if bold_prefix:
            r_bold = p.add_run(bold_prefix)
            r_bold.font.name = 'Calibri'
            r_bold.font.size = Pt(11)
            r_bold.font.bold = True
            r_bold.font.color.rgb = DARK_TEXT
        r_text = p.add_run(text)
        r_text.font.name = 'Calibri'
        r_text.font.size = Pt(11)
        r_text.font.color.rgb = DARK_TEXT
        return p

    def add_image_figure(doc, img_path, caption_text, width_inch=6.2):
        if os.path.exists(img_path):
            p_img = doc.add_paragraph()
            p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p_img.paragraph_format.space_before = Pt(10)
            p_img.paragraph_format.space_after = Pt(4)
            r_img = p_img.add_run()
            r_img.add_picture(img_path, width=Inches(width_inch))

            p_cap = doc.add_paragraph()
            p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p_cap.paragraph_format.space_before = Pt(0)
            p_cap.paragraph_format.space_after = Pt(12)
            r_cap = p_cap.add_run(caption_text)
            r_cap.font.name = 'Calibri'
            r_cap.font.size = Pt(9.5)
            r_cap.font.bold = True
            r_cap.font.color.rgb = MUTED_TEXT

    # ==================== COVER PAGE ====================
    # Logo
    logo_path = os.path.join(os.path.dirname(__file__), 'extracted', 'page_1_0_X8.png')
    if os.path.exists(logo_path):
        p_logo = doc.add_paragraph()
        p_logo.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_logo.paragraph_format.space_before = Pt(0)
        p_logo.paragraph_format.space_after = Pt(12)
        run_logo = p_logo.add_run()
        run_logo.add_picture(logo_path, width=Inches(3.2))

    # College Name
    p_inst = doc.add_paragraph()
    p_inst.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_inst.paragraph_format.space_after = Pt(2)
    r_inst = p_inst.add_run("Kolhapur Institute of Technology's\nCollege of Engineering (Autonomous), Kolhapur")
    r_inst.font.name = 'Calibri'
    r_inst.font.size = Pt(13)
    r_inst.font.bold = True
    r_inst.font.color.rgb = NAVY

    # Department
    p_dept = doc.add_paragraph()
    p_dept.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_dept.paragraph_format.space_after = Pt(20)
    r_dept = p_dept.add_run("Department of Computer Science Engineering\n(Artificial Intelligence & Machine Learning)")
    r_dept.font.name = 'Calibri'
    r_dept.font.size = Pt(11)
    r_dept.font.bold = True
    r_dept.font.color.rgb = SLATE

    # Project Title Box
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_title.paragraph_format.space_before = Pt(10)
    p_title.paragraph_format.space_after = Pt(4)
    r_title = p_title.add_run("AI INTERVIEW COACH")
    r_title.font.name = 'Calibri'
    r_title.font.size = Pt(22)
    r_title.font.bold = True
    r_title.font.color.rgb = NAVY

    p_sub = doc.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_sub.paragraph_format.space_after = Pt(16)
    r_sub = p_sub.add_run("Real-Time Mock Interview Analysis & Feedback System\n\"Practice with Data. Speak with Confidence.\"")
    r_sub.font.name = 'Calibri'
    r_sub.font.size = Pt(12)
    r_sub.font.italic = True
    r_sub.font.color.rgb = SLATE

    # Report Tag
    p_tag = doc.add_paragraph()
    p_tag.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_tag.paragraph_format.space_after = Pt(24)
    r_tag = p_tag.add_run("(Initial Project Report)\nMini Project-III (Android) [UAMIL0571]")
    r_tag.font.name = 'Calibri'
    r_tag.font.size = Pt(11.5)
    r_tag.font.bold = True
    r_tag.font.color.rgb = DARK_TEXT

    # Submitted By Header
    p_by = doc.add_paragraph()
    p_by.paragraph_format.space_after = Pt(6)
    r_by = p_by.add_run("Submitted By:")
    r_by.font.name = 'Calibri'
    r_by.font.size = Pt(11)
    r_by.font.bold = True
    r_by.font.color.rgb = NAVY

    # Team Members Table
    table_team = doc.add_table(rows=5, cols=5)
    table_team.alignment = WD_TABLE_ALIGNMENT.CENTER
    table_team.autofit = False

    col_widths = [Inches(0.7), Inches(2.2), Inches(1.5), Inches(1.0), Inches(1.1)]
    headers = ["Sr No.", "Name", "PRN", "Roll No.", "Class"]
    team_data = [
        ["1", "Yash Lad", "2425000030", "B35", "S.Y - B"],
        ["2", "Rohit Mithari", "Vacant", "B37", "S.Y - B"],
        ["3", "Rushi Parit", "Vacant", "B-36", "S.Y - B"],
        ["4", "Onkar Karande", "Vacant", "B26", "S.Y - B"]
    ]

    hdr_cells = table_team.rows[0].cells
    for i, title in enumerate(headers):
        hdr_cells[i].width = col_widths[i]
        set_cell_background(hdr_cells[i], "1E3A8A")
        set_cell_margins(hdr_cells[i], top=120, bottom=120, left=140, right=140)
        p = hdr_cells[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER if i in [0, 2, 3, 4] else WD_ALIGN_PARAGRAPH.LEFT
        r = p.add_run(title)
        r.font.name = 'Calibri'
        r.font.size = Pt(10)
        r.font.bold = True
        r.font.color.rgb = RGBColor(255, 255, 255)

    for row_idx, row_data in enumerate(team_data):
        row_cells = table_team.rows[row_idx + 1].cells
        bg_color = "F8FAFC" if row_idx % 2 == 1 else "FFFFFF"
        for col_idx, text in enumerate(row_data):
            row_cells[col_idx].width = col_widths[col_idx]
            set_cell_background(row_cells[col_idx], bg_color)
            set_cell_margins(row_cells[col_idx], top=100, bottom=100, left=140, right=140)
            p = row_cells[col_idx].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if col_idx in [0, 2, 3, 4] else WD_ALIGN_PARAGRAPH.LEFT
            r = p.add_run(text)
            r.font.name = 'Calibri'
            r.font.size = Pt(10)
            r.font.color.rgb = DARK_TEXT
            if col_idx == 1:
                r.font.bold = True

    p_ay = doc.add_paragraph()
    p_ay.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_ay.paragraph_format.space_before = Pt(30)
    p_ay.paragraph_format.space_after = Pt(0)
    r_ay = p_ay.add_run("Academic Year: 2025 – 2026")
    r_ay.font.name = 'Calibri'
    r_ay.font.size = Pt(10.5)
    r_ay.font.bold = True
    r_ay.font.color.rgb = MUTED_TEXT

    doc.add_page_break()

    # ==================== MAIN REPORT CONTENT ====================

    # 1. Abstract
    style_heading_1(doc.add_paragraph(), "1. Abstract")
    add_body_p(doc, "The AI Interview Coach is an Android mobile application engineered to democratize access to objective, data-driven mock interview practice for engineering students and job aspirants. Traditional interview preparation relies heavily on peer feedback—which is often subjective or inconsistent—or commercial coaching services that are expensive and non-scalable. The proposed solution bridges this gap by offering a privacy-first, on-device multimodal evaluation system.")
    add_body_p(doc, "By leveraging Android Jetpack CameraX for high-fps video capture, native Android SpeechRecognizer for speech-to-text processing, and computer vision models (Google ML Kit / MediaPipe) for non-verbal tracking, the application captures and analyzes practice interview responses in real time. The app evaluates speech pace (words per minute), filler word frequency (e.g., 'um', 'like'), eye contact consistency, head movement stability, and posture. Upon completing an answer, candidates receive an instant, actionable feedback report highlighting strengths, areas of improvement, and historical progress. This project delivers an accessible, repeatable, and scalable platform for building interview confidence.")

    # 2. Introduction
    style_heading_1(doc.add_paragraph(), "2. Introduction")
    add_body_p(doc, "Interview readiness is a paramount factor governing graduate employability. Despite technical competence, many candidates fail to convert interview opportunities due to poor delivery, hesitations, lack of structured answers, and weak non-verbal communication such as poor eye contact and fidgeting.")
    add_body_p(doc, "While automated technical evaluation platforms (e.g., competitive coding portals) are widely adopted, non-verbal and verbal speech evaluation remains predominantly manual. Emerging cloud-based AI tools exist, but they suffer from high subscription fees, latency, and privacy concerns associated with uploading personal facial video recordings to third-party cloud servers.")
    add_body_p(doc, "The AI Interview Coach addresses these challenges by delivering an integrated mobile experience. Operating natively on Android devices, it functions as a personal, on-demand interview trainer. The student selects an interview question from a categorized bank (Technical, HR, Behavioral), records their response, and receives instant multimodal feedback without requiring continuous internet streaming or cloud storage of raw media.")

    # 3. Problem Statement
    style_heading_1(doc.add_paragraph(), "3. Problem Statement")
    add_body_p(doc, "Students preparing for campus recruitment face significant hurdles due to the lack of structured, objective, and private mock interview tools. Specifically, the problem encompasses:")
    add_bullet_p(doc, "Peer practice lacks standard evaluation metrics, leading to biased or superficial advice.", "Subjective & Inconsistent Feedback: ")
    add_bullet_p(doc, "Candidates are unaware of their unconscious speech filler words ('um', 'ah', 'like') and speaking pace bottlenecks.", "Lack of Quantitative Speech Metrics: ")
    add_bullet_p(doc, "Self-monitoring eye contact, body alignment, and facial stress while simultaneously formulating answers is impossible without automated vision analytics.", "Unmonitored Non-Verbal Cues: ")
    add_bullet_p(doc, "Commercial human-led mock interviews are costly and difficult to schedule frequently.", "High Cost & Accessibility Barriers: ")
    add_bullet_p(doc, "Uploading video recordings of candidates to third-party servers raises serious data privacy and security concerns.", "Privacy & Latency Concerns: ")

    # 4. Existing System & Limitations
    style_heading_1(doc.add_paragraph(), "4. Existing System & Limitations")
    add_body_p(doc, "Current interview preparation methods fall into three broad categories: manual peer practice, video recording playback, and cloud-based AI platforms.")
    add_bullet_p(doc, "Students record themselves using standard phone cameras and manually re-watch videos. This is tedious, time-consuming, and lacks automated metrics.", "Standard Camera Playback: ")
    add_bullet_p(doc, "Web-based questionnaires provide static text answers but cannot evaluate verbal confidence or body language.", "Online Question Banks: ")
    add_bullet_p(doc, "Services like Big Interview or Interviewer.AI provide automated evaluation but require high-bandwidth cloud uploads, paid subscriptions, and web infrastructure.", "Cloud AI Services: ")

    add_body_p(doc, "Key Limitations of Existing Systems:", bold_prefix="", space_after=4)
    add_bullet_p(doc, "No single offline-first mobile app combines speech rate, filler words, and eye contact tracking.", "1. Lack of Integrated Mobile Analytics: ")
    add_bullet_p(doc, "Cloud video processing introduces noticeable delays before receiving feedback reports.", "2. High Processing Latency: ")
    add_bullet_p(doc, "Streaming video streams to remote servers introduces data leakage risks.", "3. Privacy Risks: ")
    add_bullet_p(doc, "No local tracking of candidate improvement over sequential practice sessions.", "4. Absence of Session Progress Tracking: ")

    # 5. Proposed System & Key Objectives
    style_heading_1(doc.add_paragraph(), "5. Proposed System & Key Objectives")
    add_body_p(doc, "The proposed AI Interview Coach is a native Android application designed to capture audio and video concurrently via CameraX. It processes speech transcripts and facial landmarks using on-device/local SDKs to deliver real-time feedback immediately post-recording.")

    style_heading_2(doc.add_paragraph(), "5.1 Key Objectives")
    add_bullet_p(doc, "Develop an intuitive Android interface featuring curated HR, Technical, and Behavioral question categories.", "1. Centralized Practice Portal: ")
    add_bullet_p(doc, "Utilize Android SpeechRecognizer to transcribe audio into text and calculate Words Per Minute (WPM) and filler word count.", "2. Automated Speech Analysis: ")
    add_bullet_p(doc, "Deploy Google ML Kit / MediaPipe Face Landmarker to monitor eye contact ratio, head orientation, and facial composure.", "3. Computer Vision Evaluation: ")
    add_bullet_p(doc, "Aggregate speech and visual parameters into a consolidated 0–100 candidate score with explicit actionable suggestions.", "4. Instant Feedback Report: ")
    add_bullet_p(doc, "Persist historic recording summaries locally using Room Database to enable progress visualization over time.", "5. Local Session History: ")

    # 6. Scope of the Project
    style_heading_1(doc.add_paragraph(), "6. Scope of the Project")
    
    table_scope = doc.add_table(rows=1, cols=2)
    table_scope.alignment = WD_TABLE_ALIGNMENT.CENTER
    table_scope.autofit = False
    
    s_cells = table_scope.rows[0].cells
    s_cells[0].width = Inches(3.2)
    s_cells[1].width = Inches(3.2)
    
    set_cell_background(s_cells[0], "1E3A8A")
    set_cell_background(s_cells[1], "334155")
    set_cell_margins(s_cells[0], top=100, bottom=100, left=140, right=140)
    set_cell_margins(s_cells[1], top=100, bottom=100, left=140, right=140)
    
    p0 = s_cells[0].paragraphs[0]
    p0.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r0 = p0.add_run("In-Scope Features")
    r0.font.bold = True
    r0.font.color.rgb = RGBColor(255, 255, 255)
    
    p1 = s_cells[1].paragraphs[0]
    p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r1 = p1.add_run("Out-of-Scope Boundaries")
    r1.font.bold = True
    r1.font.color.rgb = RGBColor(255, 255, 255)

    scope_in = [
        "Curated Question Bank (HR, Tech, Behavioral)",
        "CameraX dual video & audio recording",
        "On-device Speech-to-Text transcription",
        "Filler word detection ('um', 'like', 'uh')",
        "Eye contact & pose stability calculation",
        "Instant score report & feedback synthesis",
        "Local Room DB progress history"
    ]
    scope_out = [
        "Live human recruiter match / video calling",
        "Guaranteed job placement or hiring dispatch",
        "External biometric wearable sensor integration",
        "Multi-language speech transcription (Initial scope: English)"
    ]

    max_rows = max(len(scope_in), len(scope_out))
    for i in range(max_rows):
        row_cells = table_scope.add_row().cells
        row_cells[0].width = Inches(3.2)
        row_cells[1].width = Inches(3.2)
        bg = "F8FAFC" if i % 2 == 1 else "FFFFFF"
        set_cell_background(row_cells[0], bg)
        set_cell_background(row_cells[1], bg)
        set_cell_margins(row_cells[0], top=80, bottom=80, left=120, right=120)
        set_cell_margins(row_cells[1], top=80, bottom=80, left=120, right=120)
        
        txt_in = scope_in[i] if i < len(scope_in) else ""
        txt_out = scope_out[i] if i < len(scope_out) else ""
        
        pi = row_cells[0].paragraphs[0]
        pi.paragraph_format.space_after = Pt(0)
        ri = pi.add_run("- " + txt_in if txt_in else "")
        ri.font.name = 'Calibri'
        ri.font.size = Pt(9.5)
        ri.font.color.rgb = DARK_TEXT
        
        po = row_cells[1].paragraphs[0]
        po.paragraph_format.space_after = Pt(0)
        ro = po.add_run("- " + txt_out if txt_out else "")
        ro.font.name = 'Calibri'
        ro.font.size = Pt(9.5)
        ro.font.color.rgb = DARK_TEXT

    # 7. Stakeholders, Assumptions & Constraints
    style_heading_1(doc.add_paragraph(), "7. Stakeholders, Assumptions & Constraints")
    style_heading_2(doc.add_paragraph(), "7.1 Stakeholders")
    add_bullet_p(doc, "Students and job seekers practicing mock interviews to improve confidence and delivery.", "Primary Users: ")
    add_bullet_p(doc, "College Training & Placement Cell (TPC) tracking student interview preparedness.", "Institutional Users: ")
    add_bullet_p(doc, "Faculty and placement mentors reviewing candidate progress reports.", "Mentors & Evaluators: ")

    style_heading_2(doc.add_paragraph(), "7.2 Assumptions")
    add_bullet_p(doc, "The candidate's device runs Android 8.0 (API 26) or higher with an active front camera and microphone.")
    add_bullet_p(doc, "Practice sessions take place in moderately illuminated environments for accurate facial landmark detection.")
    add_bullet_p(doc, "The user speaks clearly in English during the recording session.")

    style_heading_2(doc.add_paragraph(), "7.3 Constraints")
    add_bullet_p(doc, "Speech recognition accuracy depends on ambient background noise levels.")
    add_bullet_p(doc, "Real-time frame processing is throttled on low-end smartphones to prevent thermal CPU scaling.")

    # 8. Project Methodology & Stages
    style_heading_1(doc.add_paragraph(), "8. Project Methodology & Stages")
    add_body_p(doc, "The project follows an Agile iterative engineering approach divided into six distinct stages:")
    add_bullet_p(doc, "Gathering user stories from students, analyzing evaluation metrics, and defining system requirements.", "Stage 1: Requirement Gathering & SRS: ")
    add_bullet_p(doc, "Designing wireframes, activity diagrams, DFDs, and Room DB schema.", "Stage 2: Architecture & UI/UX Design: ")
    add_bullet_p(doc, "Setting up Android Studio project, CameraX preview, and audio/video recorder pipeline.", "Stage 3: Camera & Recording Pipeline: ")
    add_bullet_p(doc, "Integrating SpeechRecognizer, filler detection logic, and ML Kit Face Landmarker.", "Stage 4: Multimodal ML Engine Build: ")
    add_bullet_p(doc, "Developing score aggregation rules, report view, and Room DB history persistence.", "Stage 5: Feedback & History Modules: ")
    add_bullet_p(doc, "Unit testing modules, measuring STT latency, and conducting usability trials.", "Stage 6: Testing & Optimization: ")

    # 9. System Architecture
    style_heading_1(doc.add_paragraph(), "9. System Architecture")
    add_body_p(doc, "The AI Interview Coach architecture adopts a modular client-centric design pattern:")
    add_bullet_p(doc, "Built using Jetpack ViewBinding / Fragments adhering to MVVM design principles for clean separation of concerns.", "1. Presentation Layer (UI): ")
    add_bullet_p(doc, "CameraX manages video frames; Audio Extractor provides PCM streams to the speech analyzer simultaneously.", "2. Media Stream Layer: ")
    add_bullet_p(doc, "Combines Android SpeechRecognizer (transcription & WPM/filler count) with ML Kit Face Landmarker (eye contact & posture stability).", "3. Analytics Engine: ")
    add_bullet_p(doc, "Normalizes sub-scores (Speech Pace, Filler Score, Visual Score, Content Match) into a unified 100-point index.", "4. Scoring & Rules Engine: ")
    add_bullet_p(doc, "Room Database stores local session metadata, scores, and timestamped transcripts; Firebase handles optional sync.", "5. Data Persistence Layer: ")

    # Insert Fig 1: System Architecture Image
    sys_arch_img = os.path.join(os.path.dirname(__file__), 'diagrams', 'system_architecture.png')
    add_image_figure(doc, sys_arch_img, "Fig 1. System Architecture Diagram - AI Interview Coach")

    # 9.1 Use Case Diagram
    style_heading_2(doc.add_paragraph(), "9.1 Use Case Diagram")
    add_body_p(doc, "The Use Case diagram below illustrates the interaction between the primary actor (Student/Candidate), the secondary actor (Firebase/Local DB), and the system use cases:")
    
    use_case_img = os.path.join(os.path.dirname(__file__), 'diagrams', 'use_case_diagram.png')
    add_image_figure(doc, use_case_img, "Fig 2. Use Case Diagram - System Actors & Interactions")

    # 10. Data Flow Diagrams
    style_heading_1(doc.add_paragraph(), "10. Data Flow Diagrams (DFD)")
    add_body_p(doc, "The Data Flow Diagrams depict the movement of information through the AI Interview Coach system, from initial user input and video/audio streaming to feature extraction, score aggregation, and persistent session storage.")

    style_heading_2(doc.add_paragraph(), "10.1 Level 0 (Context Diagram)")
    add_body_p(doc, "The Level 0 Context Diagram represents the high-level system boundary, showing information exchange between the Student external entity, the central application process, and the underlying data store:")

    dfd0_img = os.path.join(os.path.dirname(__file__), 'diagrams', 'dfd_level_0.png')
    add_image_figure(doc, dfd0_img, "Fig 3. Level 0 Context Data Flow Diagram")

    style_heading_2(doc.add_paragraph(), "10.2 Level 1 (Decomposed Process View)")
    add_body_p(doc, "The Level 1 DFD decomposes the system into six core sub-processes, detailing raw media streaming, parallel speech & vision analytics, score aggregation, and historical data logging:")

    dfd1_img = os.path.join(os.path.dirname(__file__), 'diagrams', 'dfd_level_1.png')
    add_image_figure(doc, dfd1_img, "Fig 4. Level 1 Decomposed Process Data Flow Diagram")

    # 11. Functional Modules
    style_heading_1(doc.add_paragraph(), "11. Functional Modules")
    add_bullet_p(doc, "Allows users to browse and filter practice questions by domain (HR, Technical, Behavioral) and difficulty level.", "Module 1 - Question Selection: ")
    add_bullet_p(doc, "Handles front-camera video preview, audio input levels, timer countdown, and smooth recording lifecycle.", "Module 2 - Recording Interface: ")
    add_bullet_p(doc, "Converts speech to text, detects filler word counts ('um', 'like', 'basically'), and computes speaking pace (WPM).", "Module 3 - Speech & Pace Analyzer: ")
    add_bullet_p(doc, "Tracks face bounding box and eye gaze vectors to calculate percentage of time candidate maintained direct eye contact.", "Module 4 - Vision Pose Engine: ")
    add_bullet_p(doc, "Displays breakdown graphs, key suggestions, transcribed text, and calculated score metric.", "Module 5 - Feedback Report Generator: ")
    add_bullet_p(doc, "Logs past practice attempts into Room DB, rendering historic score trends for self-evaluation.", "Module 6 - History & Progress Tracker: ")

    # 12. Feasibility Study
    style_heading_1(doc.add_paragraph(), "12. Feasibility Study")
    add_bullet_p(doc, "Utilizes mature Android Jetpack, ML Kit, and Speech APIs. No experimental hardware required.", "Technical Feasibility: ")
    add_bullet_p(doc, "Intuitive single-button recording workflow requires no user training.", "Operational Feasibility: ")
    add_bullet_p(doc, "Built entirely using open-source tools and on-device processing, eliminating recurring server licensing costs.", "Economic Feasibility: ")

    # 13. Technology Stack
    style_heading_1(doc.add_paragraph(), "13. Technology Stack")
    add_body_p(doc, "The technological choices below reflect a robust, well-supported mobile architecture optimized for real-time multimodal processing on Android devices.")

    # Insert Fig 5: Tech Stack Diagram
    tech_img = os.path.join(os.path.dirname(__file__), 'diagrams', 'tech_stack.png')
    add_image_figure(doc, tech_img, "Fig 5. Technology Stack Architecture Diagram")
    
    table_tech = doc.add_table(rows=8, cols=3)
    table_tech.alignment = WD_TABLE_ALIGNMENT.CENTER
    table_tech.autofit = False

    t_widths = [Inches(1.8), Inches(2.2), Inches(2.4)]
    t_headers = ["Layer", "Technology", "Purpose / Rationale"]
    t_data = [
        ["Mobile Platform", "Android SDK (Kotlin)", "Core application logic, UI lifecycle & MVVM architecture"],
        ["Media Recording", "CameraX API", "High-performance camera preview & dual video/audio capture"],
        ["Speech Processing", "Android SpeechRecognizer", "Native on-device speech-to-text conversion & timing"],
        ["Vision Analytics", "Google ML Kit / MediaPipe", "Face landmarker, eye gaze estimation & pose posture tracking"],
        ["Local Database", "Room Database (SQLite)", "Structured persistence of practice history & question bank"],
        ["Cloud Backup", "Firebase Firestore / Storage", "Optional cloud synchronization of user metrics & feedback"],
        ["Version Control", "Git & GitHub", "Source code management, branch workflows & project tracking"]
    ]

    th_cells = table_tech.rows[0].cells
    for i, title in enumerate(t_headers):
        th_cells[i].width = t_widths[i]
        set_cell_background(th_cells[i], "1E3A8A")
        set_cell_margins(th_cells[i], top=100, bottom=100, left=140, right=140)
        p = th_cells[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER if i == 0 else WD_ALIGN_PARAGRAPH.LEFT
        r = p.add_run(title)
        r.font.name = 'Calibri'
        r.font.size = Pt(10)
        r.font.bold = True
        r.font.color.rgb = RGBColor(255, 255, 255)

    for r_idx, r_vals in enumerate(t_data):
        r_cells = table_tech.rows[r_idx + 1].cells
        bg_col = "F8FAFC" if r_idx % 2 == 1 else "FFFFFF"
        for c_idx, val in enumerate(r_vals):
            r_cells[c_idx].width = t_widths[c_idx]
            set_cell_background(r_cells[c_idx], bg_col)
            set_cell_margins(r_cells[c_idx], top=90, bottom=90, left=120, right=120)
            p = r_cells[c_idx].paragraphs[0]
            r = p.add_run(val)
            r.font.name = 'Calibri'
            r.font.size = Pt(9.5)
            r.font.color.rgb = DARK_TEXT
            if c_idx == 0:
                r.font.bold = True

    # 14. Future Scope
    style_heading_1(doc.add_paragraph(), "14. Future Scope")
    add_bullet_p(doc, "Integration with Large Language Models (LLMs) to generate dynamic follow-up interview questions based on candidate answers.")
    add_bullet_p(doc, "Voice tone and emotion analysis using audio pitch processing to detect nervousness or enthusiasm.")
    add_bullet_p(doc, "Support for multi-language mock interviews to aid regional language candidates.")
    add_bullet_p(doc, "3D Avatar Recruiter simulation for immersive virtual interview environments.")

    # 15. Conclusion
    style_heading_1(doc.add_paragraph(), "15. Conclusion")
    add_body_p(doc, "The AI Interview Coach provides a practical, scalable, and privacy-first solution to a critical problem faced by engineering students preparing for recruitment drives. By combining mobile video recording, speech-to-text processing, and computer vision analytics into an intuitive Android application, the system replaces subjective peer evaluation with objective, actionable performance metrics. The system architecture, use case diagram, data flow diagrams, and tech stack diagrams presented in this initial report establish a robust engineering foundation for building an empowering self-preparation tool for candidates.")

    output_filename = os.path.join(os.path.dirname(__file__), "AI_Interview_Coach_Project_Report.docx")
    try:
        doc.save(output_filename)
        print(f"Report successfully saved to: {output_filename}")
    except PermissionError:
        fallback_filename = os.path.join(os.path.dirname(__file__), "AI_Interview_Coach_Project_Report_v2.docx")
        doc.save(fallback_filename)
        print(f"Primary file was locked. Report successfully saved to fallback path: {fallback_filename}")

if __name__ == "__main__":
    create_report()
